import json, re, operator, requests, os
import xlsxwriter
import docx
from docx.shared import Pt

#file="C:/Users/rajde/Downloads/host-Prognosis_Securities_Pvt__Ltd_Servers_to74t0.html"
#file="E:/prime/VAPT/AIC/FW__RE-VAPT___Email_Solutions___AIC/FW_ RE-VAPT _ Email Solutions _ AIC/REVAPT-Aprl_8_p63av2.html"
#file="E:/prime/VAPT/JAYSREE/server/server_byhost.html"
file=input("path:")
file=file.replace('\\','/').removeprefix('"').removeprefix("'").removesuffix('"').removesuffix("'")
out_path=file.removesuffix('.html').removesuffix('.HTML')
data=open(file,'r',encoding='utf-8').read()

def get_risk(score, flag):
	score = float(score)
	risk='Informational'
	if flag==3:
		if score<=10.0 and score>=9.0:
			risk='Critical'
		elif score<=8.9 and score>=7.0:
			risk='High'
		elif score<=6.9 and score>=4.0:
			risk='Medium'
		elif score<=3.9 and score>=0.1:
			risk='Low'
		else:
			risk='Informational'
	if flag==2:
		if score<=10.0 and score>=7.0:
			risk='High'
		elif score<=6.9 and score>=4.0:
			risk='Medium'
		elif score<=3.9 and score>= 0.0:
			risk='Low'
	return risk

def process_data(data):
	master=[]
	data=data.strip().replace('\r','').replace('\t',' ').replace('  ',' ').replace('\n','').strip().replace('style="font-size: 22px; font-weight: 700; padding: 10px 0;">','\n##$$').replace('<div class="clear"></div>','').replace('<div xmlns="" class="clear"></div>','').replace('<br>',';').replace('</br>',';')
	for d in data.split('\n'):
		if d.startswith('##$$'):
			x={}
			d1=d.replace('\n','').replace('<div xmlns="" class="details-header">','\n$$##')
			for d2 in d1.split('\n'):
				if d2.startswith('##$$'):
					temp=str(d2).split('<',1)[0].strip().removeprefix('##$$').strip()
					#x['host']=temp

				host={}
				if d2.startswith('$$##Host Information'):
					for d3 in d2.replace('\n','').replace('<tr class=""><td class="#ffffff" style=" " colspan="1">','\n').split('\n'):
						if d3.strip().startswith('IP:'):
							temp=str(d3.strip().removesuffix('</td></tr>').strip().removeprefix('IP:</td><td class="#ffffff" style=" " colspan="1">').strip().split('<',1)[0].strip())
							#print(temp)
							host['ip']=temp

						if d3.strip().startswith('Netbios Name:'):
							temp=str(d3.strip().removesuffix('</td></tr>').strip().removeprefix('Netbios Name:</td><td class="#ffffff" style=" " colspan="1">').strip().split('<',1)[0].strip())
							#print(temp)
							host['netbios']=temp

						if d3.strip().startswith('DNS Name:'):
							temp=str(d3.strip().removesuffix('</td></tr>').strip().removeprefix('DNS Name:</td><td class="#ffffff" style=" " colspan="1">').strip().split('<',1)[0].strip())
							#print(temp)
							host['dns']=temp

						if d3.strip().startswith('MAC Address:'):
							temp=str(d3.strip().removesuffix('</td></tr>').strip().removeprefix('MAC Address:</td><td class="#ffffff" style=" " colspan="1">').strip().split('<',1)[0].strip())
							#print(temp)
							host['mac']=temp

						if d3.strip().startswith('OS:'):
							temp=str(d3.strip().removesuffix('</td></tr>').strip().removeprefix('OS:</td><td class="#ffffff" style=" " colspan="1">').strip().split('<',1)[0].strip())
							#print(temp)
							host['os']=temp
					x['host'] = host

				if d2.startswith('$$##Vulnerabilities'):
					vuln=[]
					for d3 in str(d2).strip().replace('\n','').replace('onmouseover="this.style.cursor=\'pointer\'">','\n#$VULN_').split('\n'):
						if d3.startswith('#$VULN_'):
							y={}
							for d4 in str(d3).strip().replace('\n','').replace('<div class="details-header">','\n').split('\n'):
								if d4.startswith('#$VULN_'):
									temp=str(d4).split('<',1)[0].strip().split('-',1)[-1].strip()
									y['name']=str(temp)

								if d4.startswith('Synopsis<'):
									temp=str(d4).strip().removeprefix('Synopsis</div><div style="line-height: 20px; padding: 0 0 20px 0;">').removesuffix('</div>').strip()
									y['desc']=str(temp)

								if d4.startswith('Description<'):
									temp=str(d4).strip().removeprefix('Description</div><div style="line-height: 20px; padding: 0 0 20px 0;">').removesuffix('</div>').strip()
									y['imp']=str(temp)

								if d4.startswith('Solution<'):
									temp=str(d4).strip().removeprefix('Solution</div><div style="line-height: 20px; padding: 0 0 20px 0;">').removesuffix('</div>').strip()
									y['sol']=str(temp)

								if d4.startswith('CVSS v3.0 Base Score<'):
									z={}
									temp=str(d4).strip().removeprefix('CVSS v3.0 Base Score</div><div style="line-height: 20px; padding: 0 0 20px 0;">').removesuffix('</div>').strip()
									score=(temp.split(' ',1)[0].strip())
									risk=get_risk(score,3)
									z['risk']=risk
									z['score']=float(score)
									z['string']=(temp.split(' ',1)[-1].strip().removeprefix('(').removesuffix(')').strip())
									y['v3']=z

								if d4.startswith('CVSS v2.0 Base Score<'):
									z={}
									temp=str(d4).strip().removeprefix('CVSS v2.0 Base Score</div><div style="line-height: 20px; padding: 0 0 20px 0;">').removesuffix('</div>').strip()
									score=(temp.split(' ',1)[0].strip())
									risk=get_risk(score,2)
									z['risk']=risk
									z['score']=float(score)
									z['string']=(temp.split(' ',1)[-1].strip().removeprefix('(').removesuffix(')').strip())
									y['v2']=z

								if d4.startswith('See Also<'):
									link=[]
									for z in d4.replace('\n','').replace('<tr class=""><td class="#ffffff" style=" " colspan="1">','\n').split('\n'):
										if str(z).startswith('<a'):
											link.append(z.split('>',1)[-1].strip().split('<',1)[0].strip())
									link=";".join(map(str,list(set(link))))
									y['link']=link

								if d4.startswith('References<'):
									ref=[]
									for z in d4.replace('\n', '').replace('<tr class=""><td class="#ffffff" style=" " colspan="1">','\n').split('\n'):
										if re.search('</td>',z,re.IGNORECASE):
											z=(z.replace('</td><td class="#ffffff" style=" " colspan="1">','#').replace('</td></tr>','').strip().replace('</a>','').replace('</tbody></table></div>',''))
											#print(z)
											z1=z.strip().split('#',1)[0].strip()
											z2=z.strip().split('#',1)[-1].strip().split('>',1)[-1].strip().replace(':','-')
											#print('=>>',z1,z2)
											z=z1+':'+z2
											ref.append(z)
									ref=";".join(map(str,list(set(ref))))
									y['ref']=ref

								if d4.startswith('Plugin Output<'):
									port=[]
									for z in str(d4).strip().replace('<h2>','\n<h2>').replace('</h2>','</h2>\n').split('\n'):
										if z.startswith('<h2>'):
											port.append(z.strip().removeprefix('<h2>').strip().removesuffix('</h2>').strip())
									port=";".join(map(str, list(set(port))))
									y['port']=port

							#if not re.search('Nessus',y['name'],re.IGNORECASE):
								#print('===>',y['name'])
							vuln.append(y)
							#break
					x['vuln']=vuln
			if (len(x))>0:
				master.append(x)
			#break
	return master

def filter_data(data):
	vuln=[];ips=[];names=[]
	for d in data:
		x={};y={}
		ip=str(d['host']['ip'])
		try:
			y['mac']=(d['host']['mac'])
		except:
			pass
		try:
			y['os']=str(d['host']['os'])
		except:
			pass
		try:
			y['netbios']=str(d['host']['netbios'])
		except:
			pass
		try:
			y['dns']=str(d['host']['dns'])
		except:
			pass
		y['ip']=ip
		#x[ip]=y
		#ips.append(x)
		ips.append(y)
		for v in (d['vuln']):
			if 'v3' in list(v.keys()):
				v['risk']=(v['v3']['risk'])
				v['score']=(v['v3']['score'])
				v['string']=(v['v3']['string'])
			if 'v2' in list(v.keys()) and 'v3' not in list(v.keys()):
				v['risk']=(v['v2']['risk'])
				v['score']=(v['v2']['score'])
				v['string']=(v['v2']['string'])
			if 'v2' not in list(v.keys()) and 'v3' not in list(v.keys()):
				v['risk']='Informational'
				v['score']=0.0
				v['string']=''
			if not (re.search('Nessus',str(v['name']),re.IGNORECASE) or re.search('Ethernet Card Manufacturer Detection',str(v['name']),re.IGNORECASE) or re.search('Open Port Re-check',str(v['name']),re.IGNORECASE) or re.search('Host Fully Qualified Domain Name \(FQDN\) Resolution',str(v['name']),re.IGNORECASE) or re.search('Additional DNS Hostnames',str(v['name']),re.IGNORECASE) or re.search('HyperText Transfer Protocol \(HTTP\) Information',str(v['name']),re.IGNORECASE) or re.search('IMAP Service Banner Retrieval',str(v['name']),re.IGNORECASE) or re.search('DCE Services Enumeration',str(v['name']),re.IGNORECASE) or re.search('Device Type',str(v['name']),re.IGNORECASE) or re.search('Ethernet MAC Addresses',str(v['name']),re.IGNORECASE) or re.search('Session Initiation Protocol Detection',str(v['name']),re.IGNORECASE) or re.search('NetBIOS Multiple IP Address Enumeration',str(v['name']),re.IGNORECASE) or re.search('SSL / TLS Versions Supported',str(v['name']),re.IGNORECASE) or re.search('SSL Certificate Information',str(v['name']),re.IGNORECASE) or re.search('Unknown Service Detection: Banner Retrieval',str(v['name']),re.IGNORECASE) or re.search('SSL Certificate Chain Contains Certificates Expiring Soon',str(v['name']),re.IGNORECASE) or re.search('SSL Certificate Expiry - Future Expiry',str(v['name']),re.IGNORECASE) or re.search('SSL Cipher Suites Supported',str(v['name']),re.IGNORECASE) or re.search('SSL/TLS Recommended Cipher Suites',str(v['name']),re.IGNORECASE) or re.search('SSL Certificate \'commonName\' Mismatch',str(v['name']),re.IGNORECASE) or re.search('Windows NetBIOS / SMB Remote Host Information Disclosure',str(v['name']),re.IGNORECASE) or re.search('SSL Session Resume Supported',str(v['name']),re.IGNORECASE) or re.search('Terminal Services Use SSL/TLS',str(v['name']),re.IGNORECASE) or re.search('WMI Not Available',str(v['name']),re.IGNORECASE) or re.search('VNC HTTP Server Detection',str(v['name']),re.IGNORECASE) or re.search('SSL Service Requests Client Certificate',str(v['name']),re.IGNORECASE) or re.search('VNC Software Detection',str(v['name']),re.IGNORECASE) or re.search('TCP/IP Timestamps Supported',str(v['name']),re.IGNORECASE) or re.search('SSL Perfect Forward Secrecy Cipher Suites Supported',str(v['name']),re.IGNORECASE) or re.search('SSL Cipher Block Chaining Cipher Suites Supported',str(v['name']),re.IGNORECASE) or re.search('ICMP Timestamp Request Remote Date Disclosure',str(v['name']),re.IGNORECASE) or re.search('Windows Terminal Services Enabled',str(v['name']),re.IGNORECASE) or re.search('Target Credential Status by Authentication Protocol - No Credentials Provided',str(v['name']),re.IGNORECASE) or re.search('VNC Server Security Type Detection',str(v['name']),re.IGNORECASE) or re.search('Microsoft Windows SMB Service Detection',str(v['name']),re.IGNORECASE) or re.search('MSSQL Host Information in NTLM SSP',str(v['name']),re.IGNORECASE) or re.search('WS-Management Server Detection',str(v['name']),re.IGNORECASE) or re.search('Web Server No 404 Error Code Check',str(v['name']),re.IGNORECASE) or re.search('Non-compliant Strict Transport Security (STS)',str(v['name']),re.IGNORECASE) or re.search('SSL Root Certification Authority Certificate Information',str(v['name']),re.IGNORECASE) or re.search('SMTP Host Information in NTLM SSP',str(v['name']),re.IGNORECASE) or re.search('POP Server Detection',str(v['name']),re.IGNORECASE) or re.search('SMTP Server Detection',str(v['name']),re.IGNORECASE) or re.search('Non-compliant Strict Transport Security \(STS\)',str(v['name']),re.IGNORECASE) or re.search('Microsoft Outlook Web Access \(OWA\) Version Detection',str(v['name']),re.IGNORECASE) or re.search('TLS Command Support',str(v['name']),re.IGNORECASE) or re.search('Strict Transport Security (STS) Detection',str(v['name']),re.IGNORECASE) or 	re.search('Microsoft SQL Server STARTTLS Support',str(v['name']),re.IGNORECASE)	or re.search('Microsoft SQL Server TCP/IP Listener Detection',str(v['name']),re.IGNORECASE)	or re.search('Service Detection',str(v['name']),re.IGNORECASE) or re.search('Web Server Unconfigured - Default Install Page Present',str(v['name']),re.IGNORECASE) or re.search('Microsoft Windows SMB Versions Supported',str(v['name']),re.IGNORECASE) or re.search('OS Security Patch Assessment Not Available',str(v['name']),re.IGNORECASE)	or re.search('OS Identification',str(v['name']),re.IGNORECASE) or re.search('Microsoft Windows SMB NativeLanManager Remote System Information Disclosure',str(v['name']),re.IGNORECASE)	or re.search('Server Message Block \(SMB\) Protocol Version 1 Enabled',str(v['name']),re.IGNORECASE) or re.search('Traceroute Information',str(v['name']),re.IGNORECASE) or re.search('JQuery Detection',str(v['name']),re.IGNORECASE) or re.search('SSL/TLS Service Requires Client Certificate',str(v['name']),re.IGNORECASE) or re.search('Patch Report',str(v['name']),re.IGNORECASE) or re.search('Universal Plug and Play \(UPnP\) Protocol Detection',str(v['name']),re.IGNORECASE) or re.search('Web Server UPnP Detection',str(v['name']),re.IGNORECASE)):
				#re.search('',str(v['name']),re.IGNORECASE)

				#print(v['name'],v['risk'])
				names.append(v['name'])
				vuln.append(v)
	names=list(set(names))
	vuln=list({v['name']:v for v in vuln}.values())
	vuln.sort(key=operator.itemgetter('score'), reverse=True)

	for v in vuln:
		ip=[];port='';ref='';link=''
		for d in data:
			for v1 in d['vuln']:
				if str(v1['name'])==str(v['name']):
					ip.append(d['host']['ip'])
					try:
						link=(v1['link'])
					except:
						pass
					try:
						ref=(v1['ref'])
					except:
						pass
					try:
						port=(v1['port'])
					except:
						pass

		try:
			del v['v3']
		except:
			pass
		try:
			del v['v2']
		except:
			pass

		ip=list(set(ip))
		v['ip']=ip
		v['link']=link
		v['ref']=ref
		v['port']=port
		#print(v.keys())
	"""
	for i in ips:
		for v in vuln:
			if i+';' in (";".join(v['ip'])+';'):
				#print(v['risk'],v['name'])
				pass
	"""
	return vuln,names,ips

def get_count(data):
    crit=high=med=low=info=0
    for d in data:
        #print(d['risk'],d['name'])
        if str(d['risk'])=='Critical':
            crit+=1
        if str(d['risk'])=='High':
            high+=1
        if str(d['risk'])=='Medium':
            med+=1
        if str(d['risk'])=='Low':
            low+=1
        if str(d['risk'])=='Informational':
            info+=1
    total=[crit,high,med,low,info]
    return total

def create_Pie(total):
    workbook = xlsxwriter.Workbook(out_path+'_Pie.xlsx')
    worksheet = workbook.add_worksheet()
    bold = workbook.add_format({'bold': 1})
    headings = ['Risk', 'Count']
    lab = ['Critical','High','Medium','Low','Informational']
    temp = [lab, total,]
    worksheet.write_row('A1', headings, bold)
    worksheet.write_column('A2', temp[0])
    worksheet.write_column('B2', temp[1])
    chart1 = workbook.add_chart({'type': 'pie'})
    #custom=['#C00000','#FF0000','#FFC000','#92D050','#1e90ff']
    chart1.add_series({
    'name': 'Vuln Count',
    'categories': "=Sheet1!$A$2:$A$6",
    'values':     "=Sheet1!$B$2:$B$6",
    #'categories': ['Sheet1', 1, 0, 3, 0],
    #'values':     ['Sheet1', 1, 1, 3, 1],
    'data_labels':{'value':True,'category_name':True,'position':'outside_end'},#'font':'Tahoma'},#'color':custom},
    'points': [
        {'fill': {'color': '#C00000'}},
        {'fill': {'color': '#FF0000'}},
        {'fill': {'color': '#FFC000'}},
        {'fill': {'color': '#92D050'}},
        {'fill': {'color': '#00B0F0'}}]})
    chart1.set_title({'name': 'Total Vulnerability Count'})
    chart1.set_style(10)
    worksheet.insert_chart('C2', chart1, {'x_offset': 25, 'y_offset': 10})
    workbook.close()


def create_Excel(data):
	workbook = xlsxwriter.Workbook(out_path + '_Excel.xlsx')
	worksheet = workbook.add_worksheet()
	cell_format = workbook.add_format(
		{'bold': True, 'font_color': 'black', 'font_size': 11, 'bg_color': '#FFC000', 'border': 2,
		 'border_color': 'black', 'align': 'vcenter'})
	index = 1
	worksheet.write('A' + str(index), 'Sl No.', cell_format)
	worksheet.write('B' + str(index), 'Vulnerability Name', cell_format)
	worksheet.write('C' + str(index), 'Affected System', cell_format)
	worksheet.write('D' + str(index), 'Open Ports', cell_format)
	worksheet.write('E' + str(index), 'Category', cell_format)
	worksheet.write('F' + str(index), 'Classification', cell_format)
	worksheet.write('G' + str(index), 'CVSS Score', cell_format)
	worksheet.write('H' + str(index), 'CVSS Vector', cell_format)
	worksheet.write('I' + str(index), 'Description', cell_format)
	worksheet.write('J' + str(index), 'Impact', cell_format)
	worksheet.write('K' + str(index), 'Remediation', cell_format)
	worksheet.write('L' + str(index), 'Reference Link', cell_format)
	count = 1;
	for d in data:
		index += 1
		cell_format = workbook.add_format(
			{'bold': False, 'font_color': 'black', 'font_size': 11, 'border': True, 'border_color': 'black',
			 'align': 'vcenter'})
		# print(d.keys())
		worksheet.write('A' + str(index), str(count), cell_format)
		worksheet.write('B' + str(index), str(d['name']).replace('&lt;', '<'), cell_format)
		worksheet.write('C' + str(index), str("\n".join(map(str, d['ip']))), cell_format)
		worksheet.write('D' + str(index), str(d['port']).replace(';','\n'), cell_format)
		if (str(d['risk'])) == 'Critical':
			cell_format = workbook.add_format(
				{'bold': False, 'font_color': 'black', 'font_size': 11, 'border': True, 'border_color': 'black',
				 'bg_color': '#C00000', 'align': 'vcenter'})
		if (str(d['risk'])) == 'High':
			cell_format = workbook.add_format(
				{'bold': False, 'font_color': 'black', 'font_size': 11, 'border': True, 'border_color': 'black',
				 'bg_color': '#FF0000', 'align': 'vcenter'})
		if (str(d['risk'])) == 'Medium':
			cell_format = workbook.add_format(
				{'bold': False, 'font_color': 'black', 'font_size': 11, 'border': True, 'border_color': 'black',
				 'bg_color': '#FFC000', 'align': 'vcenter'})
		if (str(d['risk'])) == 'Low':
			cell_format = workbook.add_format(
				{'bold': False, 'font_color': 'black', 'font_size': 11, 'border': True, 'border_color': 'black',
				 'bg_color': '#92D050', 'align': 'vcenter'})
		if (str(d['risk'])) == 'Informational':
			cell_format = workbook.add_format(
				{'bold': False, 'font_color': 'black', 'font_size': 11, 'border': True, 'border_color': 'black',
				 'bg_color': '#00B0F0', 'align': 'vcenter'})
		worksheet.write('E' + str(index), str(d['risk']), cell_format)
		cell_format = workbook.add_format(
			{'bold': False, 'font_color': 'black', 'font_size': 11, 'border': True, 'border_color': 'black',
			 'align': 'vcenter'})
		worksheet.write('F' + str(index), str(d['ref']).replace(';', '\n'), cell_format)
		worksheet.write('G' + str(index), str(d['score']), cell_format)
		worksheet.write('H' + str(index), str(d['string']), cell_format)
		worksheet.write('I' + str(index), str(d['desc']), cell_format)
		worksheet.write('J' + str(index), str(d['imp']), cell_format)
		worksheet.write('K' + str(index), str(d['sol']), cell_format)
		worksheet.write('L' + str(index), str(d['link']).replace(';', '\n'), cell_format)
		count += 1
	# break
	workbook.close()

def create_doc(data,ips,names):
	doc=docx.Document()
	count=1
	for i in ips:
		for i1 in i:
			doc.add_heading(str(count)+'. IP: '+str(i1),1)
		table=doc.add_table(rows=1, cols=2)
		for i2 in i[i1].keys():
			row = table.add_row().cells
			row[0].text=str(i2)
			row[1].text=str(i[i1][i2])
		count+=1
		doc.add_page_break()
	doc.save(out_path+'_doc.docx')

def create_HTML(data,ips):
	def done():
		print("Created By Rajdeep\nmail:17ack312@gmail.com")
	html=''
	html=html+str('<html>\n<head><meta charset="utf-8"><meta name="viewport" content="width=device-width, initial-scale=1"><meta http-equiv="refresh" content="30" >\n<title></title>\n<style>' + requests.get('https://raw.githubusercontent.com/17ack312/myscripts/main/style_report.css').content.decode() + '</style>\n</head>\n<body>\n')
	html=html+str('<div align="center"><table width=75%>\n')
	html=html+str('<tr style="background: #F2F2F2;text-align:center;"><th rowspan="2">Sl No.</th><th rowspan="2">IP Details</th><th colspan="5">Vulnerability Type</th><th rowspan="2">Total</th></tr><tr style="background:#F2F2F2;"><th style="color: white;background: #C00000;">Critical</th><th style="color: white;background: #FF0000;">High</th><th style="color: white;background: #FFC000;">Medium</th><th style="color: white;background: #92D050;">Low</th><th style="color: white;background: #00B0F0;">Informational</th></tr>\n')

	count=1;gt=tc=th=tm=tl=ti=0
	for i in ips:
		total = crit = high = med = low = info = 0
		for d in data:
			if str(i['ip']) in d['ip']:
				if str(d['risk']) == 'Critical':
					crit += 1
				if str(d['risk']) == 'High':
					high += 1
				if str(d['risk']) == 'Medium':
					med += 1
				if str(d['risk']) == 'Low':
					low += 1
				if str(d['risk']) == 'Informational':
					info += 1
				total = crit + med + high + low + info
		html=html+str('<tr id="data_row" style="background:#F2F2F2;text-align:center;"><td id="sl">' + str(count) + '</td><td id="ip">' + str(i['ip']) + '</td><td id="critical" style="color: white;background: #C00000;">' + str(crit) + '</td><td id="high" style="color: white;background: #FF0000;">' + str(high) + '</td><td id="medium" style="color: white;background: #FFC000;">' + str(med) + '</td><td id="low" style="color: white;background: #92D050;">' + str(low) + '</td><td id="informational" style="color: white;background: #00B0F0;">' + str(info) + '</td><td id="total" style="color: white;background: #7030A0;">' + str(total) + '</td></tr>\n')
		count+=1
		gt = gt + total
		tc = tc + crit
		th = th + high
		tm = tm + med
		tl = tl + low
		ti = ti + info
	html=html+str('<tr id="data_row" style="background:#F2F2F2;text-align:center;"><td id="sl" colspan=2 style="text-align:center;">Overall Finding</td><td id="critical" style="color: white;background: #C00000;">' + str(tc) + '</td><td id="high" style="color: white;background: #FF0000;">' + str(th) + '</td><td id="medium" style="color: white;background: #FFC000;">' + str(tm) + '</td><td id="low" style="color: white;background: #92D050;">' + str(tl) + '</td><td id="informational" style="color: white;background: #00B0F0;">' + str(ti) + '</td><td id="total" style="color: white;background: #7030A0;">' + str(gt) + '</td></tr>\n')
	total = [tc, th, tm, tl, ti]
	create_Pie(total)
	html=html+str('</table></div>\n')

	html = html + str('<h1 style="page-break-after:always;"></h1>\n')

	html=html+'<ol>\n'
	for i in ips:
		html=html+('<h4 style="color:0070C0;"><li><u>IP: '+str(i['ip'])+'</u></li></h4>\n')
		sl=1;port=[]
		html=html+str('<div align="center"><table width=75%>\n')
		html = html + str('<tr style="background: #7030A0;text-align:center;color:white;font-weight:bold;"><th>Sl No.</th><th>Vulnerability Name</th><th>Risk Level</th><th>CVSS Score</th></tr>\n')
		for d in data:
			if str(i['ip']) in d['ip']:
				port.append(d['port'])
				html=html+('<tr id="v_list" style="color:black;background:#F2F2F2;text-align:center;">')
				if str(d['risk']) == 'Critical':
					html=html+('<td>' + str(sl) + '</td><td>' + str(d['name']).replace('&lt;','<') + '</td><td style="color:white;background:#C00000;">' + str(d['risk']) + '</td><td>' + str(d['score']) + '</td>')
				if str(d['risk']) == 'High':
					html=html+('<td>' + str(sl) + '</td><td>' + str(d['name']).replace('&lt;','<') + '</td><td style="color:white;background:#FF0000;">' + str(d['risk']) + '</td><td>' + str(d['score']) + '</td>')
				if str(d['risk']) == 'Medium':
					html=html+('<td>' + str(sl) + '</td><td>' + str(d['name']).replace('&lt;','<') + '</td><td style="color:white;background:#FFC000;">' + str(d['risk']) + '</td><td>' + str(d['score']) + '</td>')
				if str(d['risk']) == 'Low':
					html=html+('<td>' + str(sl) + '</td><td>' + str(d['name']).replace('&lt;','<') + '</td><td style="color:white;background:#92D050;">' + str(d['risk']) + '</td><td>' + str(d['score']) + '</td>')
				if str(d['risk']) == 'Informational':
					html=html+('<td>' + str(sl) + '</td><td>' + str(d['name']).replace('&lt;','<') + '</td><td style="color:white;background:#00B0F0;">' + str(d['risk']) + '</td><td>' + str(d['score']) + '</td>')
				html=html+('</tr>\n')
				sl+=1

		html = html + str('</table></div>\n')
		port=list(set(port))

		html=html+'<h5><u>Open Port Details</u></h5>\n'

		html=html+str('<div align="center"><table width=75%>\n')
		html = html + str('<tr style="background: #7030A0;text-align:center;color:white;font-weight:bold;"><th colspan=2>Port</th><th>State</th><th>Service</th><th>Version</th></tr>\n')
		for p in port:
			p1=str(p.strip().split('/',1)[-1].strip().split('/',1)[0].strip())
			p2=str(p.strip().split('/',1)[0].strip())
			p3='open'
			p4=str(p.strip().split('/',1)[-1].strip().split('/',1)[-1])
			html=html+str('<tr style="text-align:center;"><td>'+p1+'</td><td>'+p2+'</td><td>'+p3+'</td><td>'+p4.replace(p1,'')+'</td><td></td></tr>\n')
		html = html + str('</table></div>\n')
	html=html+'</ol>\n'

	html = html + str('<h1 style="page-break-after:always;"></h1>\n')

	count=1
	for d in data:
		ip=[]
		if str(d['risk']) == 'Critical':
			html = html + str('<div align="center"><table class="crit" id="vul_details" width="75%">\n')
		if str(d['risk']) == 'High':
			html = html + str('<div align="center"><table class="high" id="vul_details" width="75%">\n')
		if str(d['risk']) == 'Medium':
			html = html + str('<div align="center"><table class="med" id="vul_details" width="75%">\n')
		if str(d['risk']) == 'Low':
			html = html + str('<div align="center"><table class="low" id="vul_details" width="75%">\n')
		if str(d['risk']) == 'Informational':
			html = html + str('<div align="center"><table class="info" id="vul_details" width="75%">\n')

		html = html + str('<tr id="vul_name"><td colspan=2><ul type="none"><li>' + str(count) + '. Vulnerability Name: ' + str(d['name']) + '<br>\n')
		html = html + str('Vulnerability Rating: ' + str(d['risk']) + '</li></ul></td></tr>\n')

		html = html + str('<tr id="cvss"><td id="point" colspan=2>CVSS: <span id="details">' + str(d['score']) + ' ' + str(d['string']) + '</span></td></tr>\n')

		if len(str(d['ref'])) > 3:
			ref = []
			for r in d['ref'].split(';'):
				if str(r).strip().startswith('CVE:'):
					ref.append('<a href="https://nvd.nist.gov/vuln/detail/' + str(r).split(':', 1)[-1].strip() + '" target="_blank">' + str(r).split(':', 1)[-1].strip() + '</a>')
				if str(r).strip().startswith('BID:'):
					ref.append('<a href="http://www.securityfocus.com/bid/' + str(r).strip().split(':', 1)[-1].strip() + '" target="_blank">BID:' + str(r).strip().split(':', 1)[-1].strip() + '</a>')
				if str(r).strip().startswith('XREF:') and re.search('cwe', r, re.IGNORECASE):
					ref.append('<a href="https://cwe.mitre.org/data/definitions/' + str(r).strip().split(':')[
						-1].strip() + '" target="_blank">' + str(r).strip().split(':', 1)[-1].strip() + '</a>')
				if str(r).strip().startswith('XREF:') and re.search('CERT', r, re.IGNORECASE):
					ref.append('<a href="https://www.kb.cert.org/vuls/id/' + str(r).strip().split(':')[
						-1].strip() + '" target="_blank">CERT:' + str(r).strip().split(':', 1)[-1].strip() + '</a>')
			ref = ", ".join(map(str, list(set(ref))))

			html = html + str('<tr id="classification"><td id="point" colspan=2>Classification: <span id="details">' + str(ref) + '</span></td></tr>\n')

		for i in d['ip']:
			ip.append(i)
		ip=", ".join(map(str,list(set(ip))))

		html = html + str('<tr id="system"><td id="point" colspan=2>Affected Systems: <br><span id="details">' + str(ip) + '</span></td></tr>\n')

		html = html + str('<tr id="desc"><td id="point">Description: </td><td id="details">' + str(d['desc']).replace(';','<br>') + '</td></tr>\n')
		html = html + str('<tr id="impact"><td id="point">Impact: </td><td id="details">' + str(d['imp']).replace(';','<br>') + '</td></tr>\n')
		html = html + str('<tr id="rem"><td id="point">Remediation: </td><td id="details">' + str(d['sol']).replace(';', '<br>')+'\n')

		if len(str(d['link'])) > 5:
			html = html + str('<br><span><span id="ref" style="font-weight:bold;text-decoration:underline;">Reference Links: </span>\n')
			html = html + str('<ul type="circle" style="font-weight:unset">\n')
			for l in list(d['link'].split(';')):
				html = html + str('<li><a href="' + str(l) + '" target="_blank">' + str(l) + '</a></li>\n')
			html = html + str('</ul></span>\n')
		html = html + str('</td></tr>\n')

		html = html + str('<tr id="poc"><td id="point"colspan=2>Proof of Concept: </td></tr>\n')
		html = html + str('<tr id="poc"><td id="details"colspan=2></td></tr>\n')

		html = html + str('<tr id="auditee"><td id="point">Auditee\'s Response: </td><td id="details"></td></tr>\n')
		html = html + str('<tr id="remark"><td id="point">Closure Remark: </td><td id="details"></td></tr>\n')
		html = html + str('<tr id="blank"><td id="point"colspan=2></td></tr>\n')

		html = html + str('</table></div>\n')
		html = html + str('<h1 style="page-break-after:always;"></h1>\n')

		count+=1
	html=html+'</body></html>\n'

	f=open(out_path+'_Report.html','w')
	f.write(html)
	f.close()
	#done()


data=process_data(data)
data,names,ips=filter_data(data)


total=get_count(data)
#create_Pie(total)
create_Excel(data)
#create_doc(data,ips,names)
create_HTML(data,ips)



